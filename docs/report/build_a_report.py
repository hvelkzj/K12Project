from __future__ import annotations

import subprocess
import sys
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "artifacts" / "report"
FIGURE_DIR = OUTPUT_DIR / "figures"
SCREENSHOT_DIR = ROOT / "docs" / "screenshots"
REPORT_NAME = "黄章杰+黄章杰_实验报告.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17324D"
MUTED = "5E6C7B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E7F4EE"
PALE_GOLD = "FFF5D6"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 east_asia: str = "Arial Unicode MS") -> None:
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph_tokens(paragraph, *, before: float = 0, after: float = 6,
                         line: float = 1.10, keep_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_text(doc: Document, text: str, *, bold_prefix: str | None = None,
             after: float = 6, color: str = "000000"):
    paragraph = doc.add_paragraph()
    set_paragraph_tokens(paragraph, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=color)
    return paragraph


def add_bullet(doc: Document, text: str, *, level: int = 0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.5 + level * 0.25)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths_dxa: Sequence[int], *, header_fill: str = LIGHT_GRAY,
              font_size: float = 9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        set_paragraph_tokens(paragraph, after=0, line=1.05)
        run = paragraph.add_run(value)
        set_run_font(run, size=font_size, bold=True, color=INK)
    for values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            set_paragraph_tokens(paragraph, after=0, line=1.05)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
    spacer = doc.add_paragraph()
    set_paragraph_tokens(spacer, after=2)
    return table


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_tokens(paragraph, after=0, line=1.10)
    run = paragraph.add_run(f"{label}  ")
    set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(paragraph, before=3, after=8, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    return paragraph


def add_picture(doc: Document, path: Path, caption: str, *, width: float = 6.2):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(paragraph, after=0, line=1.0)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_picture_pair(doc: Document, left: tuple[Path, str], right: tuple[Path, str]):
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (path, caption) in enumerate((left, right)):
        image_p = table.cell(0, index).paragraphs[0]
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_tokens(image_p, after=0, line=1.0)
        image_p.add_run().add_picture(str(path), width=Inches(2.95))
        caption_p = table.cell(1, index).paragraphs[0]
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_tokens(caption_p, after=0, line=1.0)
        run = caption_p.add_run(caption)
        set_run_font(run, size=8.7, color=MUTED)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:val"), "nil")
                borders.append(node)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def crop_figure(source: Path, destination: Path, box: tuple[int, int, int, int]) -> Path:
    with Image.open(source) as source_image:
        source_image.convert("RGB").crop(box).save(destination, quality=94)
    return destination


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def chart_font(size: int, *, bold: bool = False):
    names = (
        "Arial Unicode.ttf",
        "Hiragino Sans GB.ttc",
        "msyhbd.ttc" if bold else "msyh.ttc",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    )
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                  *, font, fill: str, spacing: int = 6) -> None:
    bounds = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    x = box[0] + (box[2] - box[0] - width) / 2
    y = box[1] + (box[3] - box[1] - height) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, align="center", spacing=spacing)


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                *, color: str, font, fill: str = "#FFFFFF", radius: int = 20) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=color, width=3)
    centered_text(draw, box, text, font=font, fill=color)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
          label: str, *, font, color: str = "#5E6C7B") -> None:
    draw.line([start, end], fill=color, width=3)
    direction = 1 if end[0] >= start[0] else -1
    draw.polygon([
        end,
        (end[0] - direction * 15, end[1] - 8),
        (end[0] - direction * 15, end[1] + 8),
    ], fill=color)
    bounds = draw.textbbox((0, 0), label, font=font)
    width = bounds[2] - bounds[0]
    x = (start[0] + end[0]) / 2 - width / 2
    y = min(start[1], end[1]) - 28
    draw.text((x, y), label, font=font, fill=color)


def draw_architecture(path: Path) -> None:
    image = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(image)
    title_font = chart_font(42, bold=True)
    box_font = chart_font(24)
    small_font = chart_font(18)
    draw.text((900, 55), "K12 教育协同平台多端总体架构", anchor="mm",
              font=title_font, fill="#17324D")
    client_boxes = [
        ((45, 180, 280, 315), "统一入口\n注册与导航"),
        ((335, 180, 570, 315), "家长网页端"),
        ((625, 180, 860, 315), "学生网页端"),
        ((915, 180, 1150, 315), "教师网页端"),
        ((1205, 180, 1440, 315), "管理后台"),
        ((1495, 180, 1755, 315), "学生 APP /\n微信小程序"),
    ]
    for box, label in client_boxes:
        rounded_box(draw, box, label, color="#2E74B5", font=box_font)
        center_x = (box[0] + box[2]) // 2
        arrow(draw, (center_x, box[3]), (900, 440), "请求", font=small_font)
    service = (520, 440, 1280, 600)
    rounded_box(draw, service,
                "Node.js + TypeScript 统一服务\n会话、权限、业务规则、文件传输",
                color="#1F4D78", font=chart_font(29))
    shared = (190, 735, 820, 900)
    store = (980, 735, 1610, 900)
    rounded_box(draw, shared, "@k12/shared 公共契约\n六角色、实体、状态、中文标签",
                color="#2F7A65", font=chart_font(27))
    rounded_box(draw, store, "可注入内存业务仓库\n双校区种子与独立测试仓库",
                color="#7A5A00", font=chart_font(27))
    arrow(draw, (730, 600), (505, 735), "实体与状态", font=small_font)
    arrow(draw, (1070, 600), (1295, 735), "读写与联动", font=small_font)
    draw.text((900, 980), "同一学生概览契约同时服务网页、Android APP 与微信小程序",
              anchor="mm", font=chart_font(22), fill="#5E6C7B")
    image.save(path)


def draw_workflow(path: Path) -> None:
    image = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 60), "作业与批改跨端闭环", anchor="mm",
              font=chart_font(42, bold=True), fill="#17324D")
    centers = [210, 670, 1130, 1590]
    labels = ["教师端", "统一服务", "学生网页 / 移动端", "文件服务"]
    colors = ["#2E74B5", "#1F4D78", "#2F7A65", "#7A5A00"]
    for x, label, color in zip(centers, labels, colors):
        rounded_box(draw, (x - 150, 155, x + 150, 260), label,
                    color=color, font=chart_font(25))
        for y in range(280, 920, 16):
            draw.line([(x, y), (x, y + 8)], fill="#CAD2DA", width=2)
    steps = [
        (330, 0, 1, "发布作业与附件"),
        (420, 1, 2, "概览同步待完成作业"),
        (510, 2, 3, "先上传附件"),
        (600, 2, 1, "创建提交 attempt"),
        (690, 1, 0, "教师读取提交与附件"),
        (780, 0, 1, "保存分数、评语、订正要求"),
        (870, 1, 2, "回读批改结果与历史记录"),
    ]
    for y, start_i, end_i, label in steps:
        arrow(draw, (centers[start_i], y), (centers[end_i], y), label,
              font=chart_font(19), color=colors[start_i])
    image.save(path)


def draw_test_layers(path: Path) -> None:
    image = Image.new("RGB", (1800, 960), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 60), "多层测试与验收证据", anchor="mm",
              font=chart_font(42, bold=True), fill="#17324D")
    layers = [
        ((180, 690, 1620, 820), "端侧人工验收：网页、微信开发者工具、Android 运行环境", "#E8EEF5"),
        ((300, 510, 1500, 640), "真实 HTTP 与跨端集成：六角色、业务联动、权限隔离", "#DCEBE6"),
        ((420, 330, 1380, 460), "工作区自动化：客户端、状态、错误、边界与回归", "#FFF0C2"),
        ((540, 150, 1260, 280), "类型检查、Lint 与构建", "#F2F4F7"),
    ]
    for box, label, fill in layers:
        rounded_box(draw, box, label, color="#5E6C7B", fill=fill, font=chart_font(25))
    draw.text((900, 895), "332 项有效测试通过 · 500 请求 / 并发 50 / 失败 0",
              anchor="mm", font=chart_font(25), fill="#2F7A65")
    image.save(path)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(11)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("K12 教育协同平台 · A 成员实验报告")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(footer_table, [7020, 2340])
    left = footer_table.cell(0, 0).paragraphs[0]
    set_paragraph_tokens(left, after=0, line=1.0)
    run = left.add_run("高级程序设计课程项目 · 2026")
    set_run_font(run, size=8.5, color=MUTED)
    right = footer_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_tokens(right, after=0, line=1.0)
    run = right.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    right._p.append(field)
    run = right.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document, branch: str, commit: str) -> None:
    for _ in range(3):
        spacer = doc.add_paragraph()
        set_paragraph_tokens(spacer, after=24)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(kicker, after=18)
    run = kicker.add_run("高级程序设计 · 课程项目实验报告")
    set_run_font(run, size=12, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(title, after=8)
    run = title.add_run("K12 教育协同平台")
    set_run_font(run, size=30, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(subtitle, after=34)
    run = subtitle.add_run("A 成员实验报告")
    set_run_font(run, size=18, color=DARK_BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(lead, after=46, line=1.25)
    run = lead.add_run("面向家长、学生、教师与校区管理人员的\n多端家校教学协同系统")
    set_run_font(run, size=13, color=MUTED)

    meta = doc.add_table(rows=5, cols=2)
    set_table_geometry(meta, [2600, 6760])
    meta.style = "Table Grid"
    values = [
        ("姓名", "黄章杰"),
        ("项目职责", "A：公共契约、统一服务、跨端集成与质量收口"),
        ("项目基线", f"{branch} · {commit}"),
        ("完成日期", date.today().isoformat()),
        ("技术栈", "TypeScript · Vue 3 · uni-app · Node.js · npm workspaces"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for index, text in enumerate((label, value)):
            paragraph = row.cells[index].paragraphs[0]
            set_paragraph_tokens(paragraph, after=0)
            run = paragraph.add_run(text)
            set_run_font(run, size=10.5, bold=index == 0, color=INK)


def get_git_value(args: list[str], fallback: str) -> str:
    try:
        return subprocess.check_output(args, cwd=ROOT, text=True).strip()
    except (OSError, subprocess.CalledProcessError):
        return fallback


def build_report(output_path: Path) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    architecture = FIGURE_DIR / "multi-end-architecture.png"
    workflow = FIGURE_DIR / "assignment-workflow.png"
    tests = FIGURE_DIR / "test-layers.png"
    draw_architecture(architecture)
    draw_workflow(workflow)
    draw_test_layers(tests)

    branch = get_git_value(["git", "branch", "--show-current"], "develop")
    commit = get_git_value(["git", "rev-parse", "--short", "HEAD"], "最终 develop")
    ai_log_count = len(list((ROOT / "docs" / "ai-log").glob("*.md")))
    mini_screenshot = SCREENSHOT_DIR / "wechat-mini-program-home.png"
    mini_submission_screenshot = SCREENSHOT_DIR / "wechat-mini-program-submission.png"
    android_screenshot = SCREENSHOT_DIR / "android-app-home.png"
    android_project_screenshot = SCREENSHOT_DIR / "android-hbuilderx-project.png"
    apk_candidates = sorted(
        (ROOT / "artifacts").glob("**/*.apk"),
        key=lambda candidate: candidate.stat().st_mtime,
        reverse=True,
    )

    doc = Document()
    configure_document(doc)
    add_cover(doc, branch, commit)
    page_break(doc)

    add_heading(doc, "报告摘要", 1)
    add_text(doc, "本项目围绕 K12 课后教学与家校协同场景，建设统一入口、家长端、学生端、教师端、管理后台、Android APP 与微信小程序。系统以六类角色、双校区数据、统一认证、公共实体与状态为基础，完成请假、签到、作业、课件、反馈、调课、工单和账号管理等跨端联动。")
    add_text(doc, "A 成员负责项目框架、公共契约、认证与业务服务、跨端集成、PR 评审收口、统一入口与注册、附件传输、移动端、课件与考勤闭环、压力测试和最终质量验证。报告中的功能、数据、截图、测试与 Git 记录均来自最终仓库和实际执行结果。")
    add_callout(doc, "最终质量摘要", "全仓 332 项有效测试通过、0 失败；2 项真实 HTTP 测试仅因受限测试进程禁止监听回环端口而跳过。500 次概览请求在并发 50 下失败数为 0，4 个越权探针全部被拒绝。", fill=PALE_GREEN)
    add_table(doc, ["交付面", "实际结果", "证据"], [
        ("网页端", "统一入口 + 家长、学生、教师、后台四个工作区", "npm run check；真实浏览器走查"),
        ("移动端", "uni-app 一套源码构建 App 与微信小程序", "build:app；build:mp-weixin；端侧截图"),
        ("统一服务", "认证、权限、业务联动、文件传输和内存仓库", "58 项 API 有效测试"),
        ("跨平台", "npm 脚本无单系统命令；Windows 基线与 macOS 最终验证", "Windows 300 项基线；macOS 332 项最终结果"),
    ], [1700, 4260, 3400])

    add_heading(doc, "1. 功能介绍", 1)
    add_heading(doc, "1.1 业务场景与角色", 2)
    add_text(doc, "系统服务于课后培训与家校协同。家长关注绑定学生的课表、请假、通知、反馈与考勤；学生关注课程、课件、作业提交和批改；教师负责签到、作业、课件、反馈与调课；教务和系统管理员负责审批、排课、账号和工单。权限由服务端根据登录会话、授课关系、家长绑定与校区范围统一判定。")
    add_table(doc, ["角色", "主要工作", "数据范围"], [
        ("家长", "学生切换、请假、通知、反馈确认/异议、考勤", "仅绑定学生"),
        ("学生", "课程课件、作业详情、附件、提交、批改结果、考勤", "仅本人"),
        ("任课教师", "签到、作业、课件、批改、反馈、调课", "本人授课班级与课程"),
        ("班主任", "查看负责班级；教学写操作仍按本人授课", "负责班级 + 本人授课"),
        ("教务", "调课/请假审批、排课、代课、反馈工单", "所属校区"),
        ("系统管理员", "全校区概览、账号启停与管理", "机构内全部校区"),
    ], [1550, 4910, 2900])

    add_heading(doc, "1.2 多端功能模块", 2)
    add_table(doc, ["端侧", "页面/模块", "核心能力"], [
        ("统一入口", "首页、注册", "角色入口导航；家长/学生公开注册"),
        ("家长网页端", "首页、学生切换、课表、请假、通知、反馈", "绑定校验；请假与反馈闭环；中文状态"),
        ("学生网页端", "首页、课件、作业列表/详情/提交/结果", "附件下载上传；草稿；订正历史；成绩回写"),
        ("教师网页端", "课程、签到、作业/课件、批改、反馈、调课", "真实文件闭环；授课权限；失败不显示假成功"),
        ("管理后台", "看板、校区、班级、课程、排课、账号、调课、工单", "校区隔离；审批状态机；账号会话撤销"),
        ("学生 APP/小程序", "登录、首页、课件、作业、详情/提交、我的", "会话恢复；文件预览；图片提交；考勤中文展示"),
    ], [1550, 3150, 4660])

    add_heading(doc, "1.3 A 成员主要交付", 2)
    for item in (
        "建立 npm workspaces 项目框架、公共包和统一的 macOS/Windows 开发约定。",
        "实现六角色认证、会话恢复与退出，以及统一的业务 API、可注入内存仓库和权限边界。",
        "统一数字 ID、角色、状态、中文标签、StudentOverview 等公共契约，供各端复用。",
        "完成跨端作业、请假、反馈、调课、工单、账号、附件、课件与考勤业务闭环。",
        "建立统一入口与注册功能，完成移动端一套源码的 App/微信小程序构建。",
        "审核成员 PR，在评审中记录问题并在 develop 完成集成修复和回归验证。",
        "补充真实业务种子、跨端集成测试、500 请求压力测试和最终交付材料。",
    ):
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "2. AI 交互阶段概述", 1)
    add_text(doc, "AI 被用作需求梳理、代码实现建议、测试补齐和 PR 评审辅助。每次输出均由人工结合业务边界、成员分工、公共字段和实际运行结果进行判断。完整记录随代码目录中的 docs/ai-log 提交；当前共保存 %d 份 Markdown 记录。" % ai_log_count)
    add_table(doc, ["阶段", "主要提示与 AI 建议", "人工判断和执行"], [
        ("需求与字段统一", "汇总 A-E 字段与待处理问题；统一角色、状态、数字 ID", "裁决冲突字段，形成公共契约与数据库建模"),
        ("认证与共享包", "实现登录、当前用户、退出和六角色账号", "补足会话过期、角色隔离、多会话和 Windows 边界测试"),
        ("业务服务", "实现家长、学生、教师、后台接口与内存仓库", "把身份来源固定为会话，增加跨校区/跨班级拒绝和状态冲突"),
        ("PR 评审集成", "逐个检查成员 PR、提出问题并合并", "问题先写入评审；合并后在 develop 统一回归和修复"),
        ("真实业务完善", "统一入口、注册、附件、中文状态和双校区数据", "浏览器实际走查后修复英文状态、假成功和不可下载附件"),
        ("多端收口", "uni-app 学生 APP/小程序、课件与考勤、压力测试", "统一 StudentOverview；保留真实端侧限制；不把源码构建写成 APK 安装成功"),
        ("报告交付", "依据最终代码、截图、测试、PR 生成实验报告", "只引用可复查证据，Word/PDF 逐页渲染验证"),
    ], [1400, 3900, 4060], font_size=8.8)

    add_heading(doc, "2.1 代表性交互示例", 2)
    add_callout(doc, "提示词方向", "“每个成员提交的字段清单分别在分支里的 docs 中，只修改指定范围；公共字段不够先提出问题。”", fill=LIGHT_GRAY)
    add_text(doc, "判断：字段统一会影响所有端侧，因此先比较分支清单和现有共享包，再由 A 确认公共名称与状态，不在各页面复制另一套含义相同的字段。")
    add_callout(doc, "提示词方向", "“处理 PR，给出评审，提出问题并同意合并，合并后修复问题。”", fill=LIGHT_GRAY)
    add_text(doc, "判断：评审记录用于说明发现的问题；合并动作与修复动作分开，确保成员提交历史可辨识，同时用 develop 的全仓检查验证接口联动。")
    add_callout(doc, "提示词方向", "“先完成真实 APP 和微信小程序，再继续课件发布、考勤闭环、整体测试和实验报告。”", fill=LIGHT_GRAY)
    add_text(doc, "判断：选择 uni-app Vue 3 + TypeScript 共享业务源码，复用公共 StudentOverview 和已有认证/文件接口；端侧真实验收与自动化构建分别记录，避免把两者混为一谈。")

    add_heading(doc, "3. 系统设计与可视化图表", 1)
    add_heading(doc, "3.1 多端总体架构", 2)
    add_picture(doc, architecture, "图 1  多端总体架构与公共契约关系", width=6.3)
    add_text(doc, "各端使用独立页面与客户端，但通过统一服务完成认证、权限和业务联动。公共包只维护一套实体、状态和中文标签。学生网页端、Android APP 和微信小程序使用相同 StudentOverview，不需要为移动端另造字段。")

    add_heading(doc, "3.2 技术与目录设计", 2)
    add_table(doc, ["层次", "技术/目录", "设计作用"], [
        ("工作区", "npm workspaces + 根 package-lock.json", "统一安装、检查、测试与构建"),
        ("公共契约", "packages/shared", "角色、实体、状态、标签、测试账号和 StudentOverview"),
        ("统一服务", "apps/api · Node.js · TypeScript", "路由、会话、权限、业务规则、文件与仓库"),
        ("网页端", "Vue 3 + Vite + TypeScript", "统一入口与四个角色工作区"),
        ("移动端", "uni-app + Vue 3 + TypeScript", "同一源码构建 App 与微信小程序"),
        ("测试", "Node test runner + 可注入 fetch/仓库/时钟", "独立用例、错误边界、跨端流程与压力"),
    ], [1500, 3000, 4860])

    add_heading(doc, "3.3 公共学生概览", 2)
    add_text(doc, "StudentOverview 包含 student、courses、teachers、courseware、assignments、submissions 和 attendance。统一服务按登录学生过滤数据，网页端和移动端只负责中文展示与交互状态。家长概览使用家长绑定校验；教师写操作额外验证实际授课关系。")
    add_table(doc, ["字段", "含义", "主要消费者"], [
        ("student", "当前登录学生与班级/校区关系", "网页首页、移动首页/我的"),
        ("courseware", "教师发布的课件与附件", "网页课件、APP/小程序课件"),
        ("assignments", "作业内容、截止时间、迟交规则、附件", "列表与详情"),
        ("submissions", "多 attempt 提交、附件、分数、评语、订正", "状态分类、历史与结果"),
        ("attendance", "课次签到状态与记录时间", "学生与家长中文考勤"),
    ], [1700, 4160, 3500])

    add_heading(doc, "3.4 权限判定", 2)
    add_table(doc, ["操作", "允许范围", "服务端拒绝条件"], [
        ("家长读取学生", "ParentStudentBinding 中的学生", "未绑定学生 403"),
        ("学生读取与提交", "当前会话学生本人", "伪造用户 ID 不生效；跨学生拒绝"),
        ("教师教学写操作", "本人实际授课课次、班级和课程", "仅班主任关系不足以写其他教师课次"),
        ("教务审批和排课", "所属校区", "跨校区 403"),
        ("系统管理员", "机构全部校区", "业务状态冲突仍返回 409"),
        ("附件读取", "附件关联业务 + 当前角色范围", "伪造归属或越权下载拒绝"),
    ], [1900, 3800, 3660])

    add_heading(doc, "3.5 作业与批改流程", 2)
    add_picture(doc, workflow, "图 2  APP/小程序提交、教师批改与移动端回读流程", width=6.3)
    add_text(doc, "附件必须先上传成功，再创建 Submission；同一作业的订正会生成更大的 attempt，旧提交继续保留。教师批改写入统一仓库后，学生网页和移动端重新加载概览即可显示分数、评语和订正状态。")

    add_heading(doc, "3.6 数据仓库设计", 2)
    add_text(doc, "当前课程项目运行时使用可注入的进程内业务仓库。服务启动时载入双校区与十三个公共账号的种子数据；服务重启后恢复初始状态。测试为每个场景创建独立仓库，并可注入时钟固定服务端时间。仓库接口与数据库迁移建模分离，后续若接入 PostgreSQL，可保持上层路由与公共契约不变。")
    add_callout(doc, "边界说明", "当前数据和附件字节不持久化到 PostgreSQL；这是本课程项目的明确范围。报告不把迁移建模描述成已上线数据库。", fill=PALE_GOLD)

    add_heading(doc, "4. 测试、调试与结果", 1)
    add_heading(doc, "4.1 测试分层", 2)
    add_picture(doc, tests, "图 3  自动化、集成、压力与端侧验收层次", width=6.2)
    add_table(doc, ["工作区", "有效通过", "失败", "重点覆盖"], [
        ("shared", "16", "0", "工作区约束、角色、账号、状态中文标签"),
        ("API", "58", "0", "认证、权限、跨端流程、附件、课件、考勤、注册"),
        ("统一入口", "7", "0", "注册字段、错误、角色跳转"),
        ("家长端", "39", "0", "绑定、请假、通知、反馈、错误状态"),
        ("学生网页端", "69", "0", "草稿、筛选、提交、附件、订正、结果"),
        ("教师端", "53", "0", "认证、签到、作业、课件、批改、反馈、调课"),
        ("后台", "57", "0", "校区范围、审批、代课、工单、排课、账号"),
        ("学生移动端", "33", "0", "登录、真机连接、超时、中国时区、相机、课件、作业、附件、考勤"),
        ("合计", "332", "0", "另有 API 2 项在受限进程按环境跳过"),
    ], [1750, 1100, 900, 5610], font_size=8.7)

    add_heading(doc, "4.2 自动化执行", 2)
    add_table(doc, ["环境", "命令", "结果"], [
        ("macOS 最终分支", "npm ci", "依赖可由根锁文件安装"),
        ("macOS 最终分支", "npm run check", "Lint、类型检查、332 项有效测试和全部构建通过"),
        ("macOS 最终分支", "npm run build:app", "生成 uni-app App 运行资源"),
        ("macOS 最终分支", "npm run build:mp-weixin", "生成可导入微信开发者工具的目录"),
        ("Windows 11 基线 ce25b32", "npm ci / npm run check / npm run dev", "300 项测试、构建、六服务与十三账号验证通过"),
    ], [1700, 3270, 4390])
    add_text(doc, "Windows 记录来自 2026-08-24 的 develop 基线。移动端新增后的 Windows 复验应使用相同根脚本执行；报告只把已经执行过的基线写为通过。")

    add_heading(doc, "4.3 压力测试", 2)
    add_table(doc, ["指标", "结果", "说明"], [
        ("业务请求", "500", "家长、学生、教师和后台概览"),
        ("最高并发", "50", "六类账号先登录并复用会话"),
        ("吞吐量", "5760.16 请求/秒", "当前 macOS 测试设备记录"),
        ("平均 / P50", "8.35 / 3.96 ms", "真实 HTTP 进程内仓库"),
        ("P95 / 最大", "48.33 / 81.59 ms", "不设依赖设备性能的硬门槛"),
        ("失败数", "0", "服务未中断，无非预期响应"),
        ("越权探针", "4/4 被拒绝", "跨角色入口均返回 403"),
    ], [1900, 2350, 5110])

    add_heading(doc, "4.4 关键问题与调试", 2)
    add_table(doc, ["现象", "定位", "修复与验证"], [
        ("页面显示 PENDING_PARENT、SUBMITTED 等英文状态", "各成员页面直接渲染公共状态码", "公共 labels 统一中文映射；共享层测试覆盖所有面向业务用户的状态"),
        ("作业附件无法下载，教师无法查看学生提交", "页面只显示元数据，缺少真实文件字节和权限接口", "增加上传/下载、归属校验和教师/学生双向附件闭环测试"),
        ("移动文件上传后大小与格式判断失真", "端侧读取 base64，而服务端原先按字符串字节校验", "增加传输标识，服务端严格解码后按原始字节执行 10 MB 与 MIME 规则"),
        ("微信小程序预填账号登录返回 401", "登录页样例密码与公共账号密码不一致", "统一为 K12Demo123!，补客户端测试并重新构建两端"),
        ("Windows Node 24 启动出现 DEP0190", "统一启动脚本使用 shell: true", "改为当前 Node 进程调用 npm_execpath；Windows 重新启动验证"),
        ("Android 真机登录长时间等待后网络失败", "模拟器专用 10.0.2.2 不能访问电脑，且请求未显式超时", "登录页保存局域网地址；API 输出手机地址；10 秒超时并补 3 项回归测试"),
        ("移动端时间显示英文，拍照提示缺少 Camera 模块", "Android WebView 未稳定采用中文区域；旧包只有权限声明而未启用原生模块", "固定按 UTC+8 输出中文年月日；启用 Camera 模块，拆分拍照/相册并补 8 项回归测试"),
    ], [2800, 2700, 3860], font_size=8.0)

    add_heading(doc, "5. 使用示例", 1)
    add_heading(doc, "5.1 统一入口与注册", 2)
    portal = SCREENSHOT_DIR / "portal-home.jpg"
    registration = SCREENSHOT_DIR / "portal-registration.jpg"
    if portal.exists() and registration.exists():
        add_picture_pair(doc, (portal, "图 4  统一入口面向家校用户介绍各端"),
                         (registration, "图 5  家长/学生公开注册"))
    add_number(doc, "访问统一入口，按身份选择家长、学生、教师或管理工作区。")
    add_number(doc, "家长或学生可填写用户名、姓名、密码和角色完成注册。")
    add_number(doc, "注册成功后进入对应登录页；教师和管理员账号由校区后台维护。")

    add_heading(doc, "5.2 教师发布课件与学生读取", 2)
    teacher_courseware = SCREENSHOT_DIR / "teacher-courseware-publish.png"
    student_courseware = SCREENSHOT_DIR / "student-courseware-sync.png"
    if teacher_courseware.exists() and student_courseware.exists():
        add_picture_pair(doc,
                         (teacher_courseware, "图 6  教师选择课次并发布课件"),
                         (student_courseware, "图 7  学生端同步读取课件"))
    add_text(doc, "教师选择本人授课课次，填写标题、说明并上传附件。服务端校验授课关系、附件所有者和重复标题后保存，学生网页、APP 与小程序通过同一学生概览读取。")

    add_heading(doc, "5.3 学生与家长查看考勤", 2)
    student_attendance = SCREENSHOT_DIR / "student-attendance.png"
    parent_attendance = SCREENSHOT_DIR / "parent-attendance.png"
    if student_attendance.exists() and parent_attendance.exists():
        add_picture_pair(doc,
                         (student_attendance, "图 8  学生首页显示本人考勤统计"),
                         (parent_attendance, "图 9  家长课表显示绑定学生考勤"))
    add_text(doc, "教师保存签到后，学生读取本人记录，家长读取绑定学生记录。页面统一显示“已出勤、迟到、请假、缺勤、未签到”，不向业务用户暴露内部状态码。")

    add_heading(doc, "5.4 APP 与微信小程序", 2)
    mobile_pairs: list[tuple[Path, str]] = []
    if mini_screenshot.exists():
        mini_home_figure = crop_figure(
            mini_screenshot, FIGURE_DIR / "mini-program-home-crop.jpg", (720, 24, 1125, 744)
        )
        mobile_pairs.append((mini_home_figure, "图 10  小程序首页、作业与考勤"))
    if mini_submission_screenshot.exists():
        mini_submission_figure = crop_figure(
            mini_submission_screenshot, FIGURE_DIR / "mini-program-submission-crop.jpg", (720, 24, 1125, 744)
        )
        mobile_pairs.append((mini_submission_figure, "图 11  小程序提交成功与记录回写"))
    if android_screenshot.exists():
        mobile_pairs.append((android_screenshot, "图 12  Android 测试应用首页"))
    if len(mobile_pairs) == 2:
        add_picture_pair(doc, mobile_pairs[0], mobile_pairs[1])
    elif len(mobile_pairs) == 1:
        add_picture(doc, mobile_pairs[0][0], mobile_pairs[0][1], width=3.1)
    if not android_screenshot.exists() and android_project_screenshot.exists():
        add_picture(doc, android_project_screenshot, "图 12  HBuilderX 5.24 已识别 student-mobile 工程", width=5.7)
    add_table(doc, ["能力", "APP", "微信小程序"], [
        ("登录与会话", "学生账号、本地保存、启动恢复", "学生账号、本地保存、启动恢复"),
        ("课件附件", "下载后调用系统打开", "下载后预览/打开"),
        ("提交附件", "拍照或选择 JPG/PNG", "图片或微信会话 PDF/DOCX/JPG/PNG"),
        ("错误处理", "401/403/409/422/网络错误中文提示", "同一客户端与页面状态"),
        ("服务地址", "模拟器默认 10.0.2.2；真机登录页保存电脑地址", "开发配置使用本机服务；发布需 HTTPS 域名"),
    ], [1750, 3650, 3960])
    if apk_candidates and android_screenshot.exists():
        add_callout(doc, "Android 验收", f"已生成测试 APK：{apk_candidates[0].name}，并保存安装/运行证据。", fill=PALE_GREEN)
    elif apk_candidates:
        add_callout(doc, "Android APK", f"已使用正式 AppID、云端证书和快速安心模式免费生成测试 APK：{apk_candidates[0].name}。0.1.1 增加局域网连接设置、健康检查和 10 秒超时；0.1.2 固定中国时区中文时间并启用 Camera 原生模块。新版完整业务流程仍待真机复验。", fill=PALE_GREEN)
    else:
        add_callout(doc, "Android 验收状态", "App 运行资源与自动化构建已经通过，HBuilderX 5.24 已识别工程并安装真机运行插件。云打包继续操作时要求 DCloud 登录，本机也未连接 Android 真机/模拟器，因此未生成 APK，安装和原生业务流程不记为通过。", fill=PALE_GOLD)

    page_break(doc)
    add_heading(doc, "6. 团队合作、接口衔接与 Git", 1)
    add_heading(doc, "6.1 分工与协同", 2)
    add_table(doc, ["成员", "主要职责", "与 A 的接口衔接"], [
        ("A", "架构、公共契约、认证/业务服务、集成测试与收口", "审核公共变化并为各端提供稳定契约"),
        ("B", "家长端", "使用家长绑定、请假、反馈、通知和考勤接口"),
        ("C", "学生网页端", "使用学生概览、提交与文件接口"),
        ("D", "教师端", "使用签到、作业、课件、批改、反馈和调课接口"),
        ("E", "管理后台", "使用概览、审批、排课、账号和工单接口"),
    ], [900, 3580, 4880])
    add_text(doc, "公共契约由 A 统一审核，成员在各自工作区实现页面。PR 审核先记录接口、状态、权限或测试问题，阻塞问题解决后合并到 develop；合并后再次运行根检查，防止单端通过但跨端联动失效。")

    add_heading(doc, "6.2 A 的主要合并记录", 2)
    add_table(doc, ["PR", "分支/主题", "主要结果"], [
        ("#2", "A 初始架构", "项目工作区、文档与协作基础"),
        ("#4", "feature/A-auth", "公共类型、账号与认证会话"),
        ("#7", "feature/A-shared-integration", "前端共享依赖统一"),
        ("#10", "feature/A-business-api-round-1", "四端业务 API 与权限联动"),
        ("#16", "feature/A-management-api-round-2", "请假、排课与账号管理接口"),
        ("#22", "feature/A-cross-end-round-3", "跨端流程与独立仓库集成测试"),
        ("#32", "fix/A-final-integration-round-4", "最终认证、脚本与 Windows 兼容收口"),
        ("#33", "feature/A-portal-registration", "统一入口、注册和真实业务验证"),
        ("#34", "feature/A-business-ui-completion", "真实附件、中文状态与业务页面完善"),
    ], [900, 3420, 5040], font_size=8.7)

    add_heading(doc, "6.3 提交与评审规范", 2)
    for item in (
        "develop 作为集成分支，功能与修复在 feature/A-* 或 fix/A-* 分支完成。",
        "提交按 feat、fix、test、docs 分类，使功能、问题修复和验证记录可辨识。",
        "公共角色、状态、API 契约和锁文件影响在 PR 描述中说明。",
        "评审问题进入 docs/ai-log 与交付记录，合并后用 npm run check 和跨端流程再次验证。",
        "项目当前保留多名成员独立提交历史，Git 记录可追溯具体协作过程。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "7. 项目规则与完整交互记录", 1)
    add_heading(doc, "7.1 项目规则", 2)
    add_table(doc, ["规则", "执行方式"], [
        ("跨平台", "Node.js + npm 脚本；不在公共脚本使用 cp/rm/export/set 等单系统命令"),
        ("依赖", "Node.js 22.12+、npm 10+、根 package-lock.json；拉取后 npm ci"),
        ("公共契约", "数据库迁移、API、角色与状态由 A 审核后合并"),
        ("测试", "合并前运行 lint、类型检查、单元/集成测试和构建"),
        ("AI 使用", "记录目标、完整提示、回复摘要、人工判断、修改、测试与下一轮问题"),
        ("安全", "不提交环境密钥、证书、编辑器配置、缓存和 node_modules"),
    ], [1900, 7460])

    add_heading(doc, "7.2 随提交包保留的过程材料", 2)
    add_table(doc, ["材料", "位置", "用途"], [
        ("协作规范", "AGENTS.md", "分支、提交、公共字段、跨平台与 AI 规则"),
        ("AI 交互", "docs/ai-log/*.md", f"{ai_log_count} 份阶段记录与 PR 评审记录"),
        ("任务记录", "docs/tasks/", "按日期、轮次与成员保存执行边界"),
        ("公共契约", "docs/api-field-contract.md", "接口字段、身份范围和错误结构"),
        ("业务规则", "docs/business-rules.md", "状态机、权限与跨端联动"),
        ("测试证据", "docs/windows-test-2026-08-24.md、docs/load-test-2026-08-25.md", "环境、命令、指标和修复"),
        ("页面证据", "docs/screenshots/", "统一入口、注册、课件、考勤与端侧页面"),
    ], [1700, 4050, 3610], font_size=8.9)
    add_text(doc, "代码目录保留完整项目文件、测试、文档和锁文件。打开提交 ZIP 后可直接看到实验报告 Word 和代码目录；代码目录排除 .git、node_modules、缓存、覆盖率、临时文件、环境密钥与编辑器配置。")

    add_callout(doc, "结论与复查", "项目已形成统一入口、公共契约、服务端权限、跨端业务联动、真实附件和中文业务状态。A 的交付覆盖架构、公共服务、成员接口衔接、移动端扩展与质量验证。运行 npm ci 后，可用 npm run check 验证全部工作区，用 npm run test:load 复现压力测试，并用 build:app 与 build:mp-weixin 重新生成移动端目录。", fill=PALE_GREEN)

    doc.core_properties.title = "K12 教育协同平台 A 成员实验报告"
    doc.core_properties.subject = "高级程序设计课程项目"
    doc.core_properties.author = "黄章杰"
    doc.core_properties.keywords = "K12, Vue, TypeScript, uni-app, Android, 微信小程序, 实验报告"
    doc.save(output_path)


def main() -> None:
    output = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else OUTPUT_DIR / REPORT_NAME
    output.parent.mkdir(parents=True, exist_ok=True)
    build_report(output)
    print(output)


if __name__ == "__main__":
    main()
